import os
import re
import yt_dlp
print("yt-dlp version:", yt_dlp.version.__version__)
import whisper
import google.generativeai as genai
from docx import Document
import uuid

# In-memory store for task progress
tasks = {}
download_paths = {}

def update_status(task_id, message, error=False, completed=False):
    tasks[task_id] = {
        "message": message,
        "error": error,
        "completed": completed
    }

def get_task_status(task_id):
    return tasks.get(task_id, {"message": "Initializing...", "error": False, "completed": False})

def get_download_path(task_id):
    return download_paths.get(task_id)

def process_video(youtube_url, task_id):
    update_status(task_id, "Extracting audio from YouTube...")
    audio_file = None
    try:
        # Setup API Key inside the worker
        API_KEY = os.environ.get("GEMINI_API_KEY")
        if not API_KEY:
            raise Exception("GEMINI_API_KEY is not set on the server.")
        genai.configure(api_key=API_KEY)
        model = genai.GenerativeModel('gemini-3.5-flash')

        # Download Audio
        uid = str(uuid.uuid4())[:8]
        ydl_opts = {
            'format': 'bestaudio/best',
            'outtmpl': f'temp_audio_{uid}.%(ext)s',
            'postprocessors': [{
                'key': 'FFmpegExtractAudio',
                'preferredcodec': 'mp3',
                'preferredquality': '192',
            }],
            'quiet': True,
            'no_warnings': True
        }
        
        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            info = ydl.extract_info(youtube_url, download=True)
            title = info.get('title', 'Unknown Title')
        
        audio_file = f'temp_audio_{uid}.mp3'

        # Transcribe
        update_status(task_id, "Transcribing audio (this may take a minute)...")
        whisper_model = whisper.load_model("base")
        result = whisper_model.transcribe(audio_file)
        transcript = result["text"]

        # Generate Notes
        update_status(task_id, "Generating structured notes using Gemini...")
        prompt = f"""
        Please read the following transcript from a video lecture and create structured, easy-to-understand notes.
        Use simple English, avoiding overly advanced vocabulary. 
        Format the notes with clear headings, bullet points, and highlight the key takeaways.

        Transcript:
        {transcript}
        """
        response = model.generate_content(prompt)
        notes = response.text

        # Save to Word
        update_status(task_id, "Saving document...")
        safe_title = re.sub(r'[\\/*?:"<>|]', "", title)
        doc_filename = f"{safe_title}_Notes.docx"
        doc_path = os.path.join(os.getcwd(), doc_filename)
        
        doc = Document()
        doc.add_heading(title, 0)
        doc.add_paragraph(notes)
        doc.save(doc_path)
        
        download_paths[task_id] = doc_path
        update_status(task_id, "Done!", completed=True)

    except Exception as e:
        update_status(task_id, f"Error: {str(e)}", error=True)
    finally:
        if audio_file and os.path.exists(audio_file):
            os.remove(audio_file)
